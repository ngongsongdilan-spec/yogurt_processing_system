from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x3D, 0x2C, 0x1A)

def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def img_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('─' * 50)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r.font.size = Pt(8)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'[ {caption} ]')
    r2.bold = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('─' * 50)
    r3.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r3.font.size = Pt(8)
    doc.add_paragraph()

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# ══════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('System Implementation')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0xE8, 0x6B, 0x8A)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run('Yogurt Processing System — Frontend, Backend & Database Integration')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
r.italic = True

doc.add_paragraph()
doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info_p.add_run('PostgreSQL • Django 6.0 • Raw SQL • Vanilla JS SPA')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════
heading('1. Implementation Overview', level=1)
para(
    'The Yogurt Processing System is implemented as a three-tier web application. The PostgreSQL '
    'database serves as the persistence layer with all business rules (constraints, triggers, '
    'referential integrity) enforced at the database level. Django acts as a thin REST API layer '
    'executing raw SQL queries. The frontend is a single-page application built with vanilla '
    'HTML, CSS, and JavaScript, communicating with the backend via the Fetch API.'
)

para('The system covers the complete production lifecycle:', size=11)
add_bullet('Raw material procurement through purchase orders')
add_bullet('Production batch creation and tracking (Planned → InProgress → Completed)')
add_bullet('Quality control inspections with pass/fail recording')
add_bullet('Inventory management for raw materials and finished goods')
add_bullet('Sales order fulfillment and delivery tracking')
add_bullet('Equipment maintenance scheduling')
add_bullet('Automated audit logging of all data changes')

img_placeholder('Figure 1: System Architecture Diagram — Browser ↔ Django (REST API) ↔ PostgreSQL')

# ══════════════════════════════════════════════════════════════
# 2. DATABASE CONNECTION
# ══════════════════════════════════════════════════════════════
heading('2. Database Connection & Configuration', level=1)

para(
    'The application connects to a PostgreSQL 17 database named yogurt hosted on localhost:5432. '
    'The connection is configured through Django\'s settings.py using the psycopg2 adapter. '
    'No Django ORM models are defined; all data access is performed via django.db.connection.cursor() '
    'with raw SQL statements.'
)

heading('2.1 Connection Configuration', level=2)
add_table(
    ['Parameter', 'Value'],
    [
        ['Database', 'yogurt'],
        ['Host', 'localhost'],
        ['Port', '5432'],
        ['User', 'postgres'],
        ['Engine', 'django.db.backends.postgresql'],
        ['Adapter', 'psycopg2 2.9.12'],
        ['Schema', 'public'],
        ['Auth apps', 'None — custom auth via raw SQL queries'],
    ]
)

para(
    'Django\'s default database configuration points directly to the yogurt database with no '
    'migration-generating apps installed. All 26 tables exist exclusively in PostgreSQL and were '
    'created through direct SQL execution rather than Django migrations.'
)

img_placeholder('Figure 2: Database Connection — settings.py Configuration')

# ══════════════════════════════════════════════════════════════
# 3. BACKEND (DJANGO REST API)
# ══════════════════════════════════════════════════════════════
heading('3. Backend — Django REST API', level=1)

para(
    'The backend is a Django 6.0 application that exposes 13 REST API endpoints. Every endpoint '
    'uses raw PostgreSQL queries executed through Django\'s connection cursor — no ORM models, '
    'no querysets, no model serializers. This approach demonstrates direct database interaction '
    'and the ability to leverage PostgreSQL-specific features (JSONB, window functions, '
    'generated columns).'
)

para(
    'Each view function receives the HTTP request, executes one or more raw SQL statements, '
    'and returns a JSON response. Error handling captures DatabaseError exceptions from '
    'PostgreSQL and returns the database error message directly to the client — useful for '
    'demonstrating constraint and trigger enforcement.'
)

heading('3.1 API Endpoints', level=2)
add_table(
    ['Method', 'Endpoint', 'Description'],
    [
        ['POST', '/api/login/', 'Authenticate user, return role + name'],
        ['POST', '/api/signup/', 'Register new user account'],
        ['GET', '/api/dashboard/', '8 factory KPIs (batches, stock, QC, machines)'],
        ['GET', '/api/inventory/', 'Raw materials + finished goods stock'],
        ['GET', '/api/batches/', 'All production batch records'],
        ['GET', '/api/qc/', 'Quality control inspection history'],
        ['POST', '/api/qc/', 'Submit new QC inspection result'],
        ['GET', '/api/machines/', 'Machine list with maintenance schedule'],
        ['GET', '/api/audit/', 'Audit log — all data changes (84+ entries)'],
        ['GET', '/api/purchasing/', 'Purchase orders with line items'],
        ['GET', '/api/sales/', 'Sales orders with line items'],
        ['POST', '/api/trigger-test/', 'Test database constraints and triggers'],
        ['POST', '/api/sql-explorer/', 'Execute ad-hoc SQL queries'],
    ]
)

heading('3.2 Raw SQL Approach', level=2)
para(
    'All views execute SQL directly. For example, the dashboard endpoint runs multiple queries '
    'to compute the 8 KPI metrics:'
)
para(
    'SELECT COUNT(*) FROM production_batch WHERE status IN (\'InProgress\', \'Planned\');\n'
    'SELECT COALESCE(SUM(quantity_on_hand), 0) FROM inventory;\n'
    'SELECT COUNT(*) FROM quality_control WHERE status = \'Fail\';\n'
    'SELECT COUNT(*) FROM audit_log WHERE changed_at >= NOW() - INTERVAL \'24 hours\';',
    size=9
)

heading('3.3 Audit Trail Implementation', level=2)
para(
    'The audit system uses a PostgreSQL trigger function (fn_audit_trigger) attached to 8 tables. '
    'On INSERT, UPDATE, or DELETE, the trigger captures full JSONB snapshots of row data and writes '
    'to the audit_log table. The acting user is identified via the PostgreSQL session variable '
    'app.username, falling back to "system" for automated processes. A COALESCE/NULLIF pattern '
    'prevents NULL overrides when the session variable is not set.'
)

img_placeholder('Figure 3: Backend — API Response Example from Dashboard Endpoint')

# ══════════════════════════════════════════════════════════════
# 4. FRONTEND
# ══════════════════════════════════════════════════════════════
heading('4. Frontend — Single-Page Application', level=1)

para(
    'The frontend is implemented entirely with vanilla HTML5, CSS3, and JavaScript. No frameworks '
    '(React, Vue, Angular) or UI libraries (Bootstrap, Tailwind) are used. The application is '
    'served by Django\'s template engine as a single index.html file with external CSS and JS.'
)

heading('4.1 Design System', level=2)
para(
    'The user interface features a creamy yogurt-inspired light theme with a warm, inviting palette:'
)
add_table(
    ['Element', 'Color', 'Purpose'],
    [
        ['Background', '#FDF6EC', 'Warm cream base'],
        ['Primary', '#E86B8A', 'Strawberry pink — brand accent, buttons, active states'],
        ['Secondary', '#E8A84C', 'Honey gold — secondary accents'],
        ['Body text', '#3D2C1A', 'Warm brown — readable, soft on eyes'],
        ['Muted text', '#8B7355', 'Muted brown — labels, descriptions'],
        ['Cards', 'rgba(255,248,240,0.75)', 'Creamy translucent glass panels with blur'],
    ]
)
para(
    'Typography uses three Google Fonts: Plus Jakarta Sans for body text, Space Grotesk for '
    'headings and metric values, and JetBrains Mono for data tables and code blocks.'
)

heading('4.2 Layout & Navigation', level=2)
para(
    'The interface is divided into a fixed sidebar and a main content area. The sidebar contains '
    '10 navigation items corresponding to functional modules. Clicking a tab switches the main '
    'content view without page reload. The sidebar also displays the authenticated user and a '
    'live connection status indicator.'
)

add_table(
    ['Tab', 'Content'],
    [
        ['Dashboard', '8 KPI metric cards, recent batches table, recent QC table'],
        ['Inventory', 'Raw materials stock + finished goods stock tables'],
        ['Production', 'All batch records with status, recipe, creator'],
        ['Quality Control', 'Inspection history + form to submit new QC results'],
        ['Machines', 'Equipment list with type, status, maintenance count, next due date'],
        ['Audit Trail', 'Full change log with table name, operation, user, timestamp'],
        ['Security', '4 test buttons to verify triggers and constraints'],
        ['Purchasing', 'Purchase orders with supplier, total, line count'],
        ['Sales', 'Sales orders with customer, total, line count'],
        ['SQL Explorer', 'Ad-hoc query input with results grid'],
    ]
)

heading('4.3 Authentication Flow', level=2)
para(
    'Users authenticate through a modal overlay. Credentials are stored in sessionStorage, '
    'persisting across page refreshes during the session. The logout button clears the session '
    'and returns to the login screen. Authentication is handled client-side with API calls '
    'to the backend. No Django auth middleware is used.'
)

img_placeholder('Figure 4a: Frontend — Login / Signup Auth Overlay')

heading('4.4 Data Presentation', level=2)
para(
    'All data from the backend is rendered into HTML tables. Status values (batch status, QC '
    'results, machine status) are color-coded with badges. The dashboard displays 8 metric cards '
    'showing key factory KPIs. Toast notifications provide feedback for all user actions.'
)

img_placeholder('Figure 4b: Frontend — Dashboard Tab with Metric Cards and Recent Batches Table')

heading('4.5 Security Sandbox', level=2)
para(
    'A dedicated Security tab provides four interactive tests that demonstrate database-level '
    'integrity enforcement. Each test attempts a prohibited operation and displays the PostgreSQL '
    'error message returned by the backend, confirming that constraints and triggers are active.'
)

add_table(
    ['Test', 'Attempted Operation', 'Expected Protection'],
    [
        ['Completed Batch Lock', 'UPDATE a completed batch', 'BEFORE UPDATE trigger blocks modification'],
        ['QC Approval Lock', 'UPDATE a passed QC record', 'BEFORE UPDATE trigger blocks changes'],
        ['Check Constraint', 'INSERT with negative quantity', 'CHECK constraint rejects invalid data'],
        ['Unique Violation', 'INSERT duplicate username', 'UNIQUE constraint prevents duplicates'],
    ]
)

img_placeholder('Figure 4c: Frontend — Security Sandbox with Constraint Test Results')

heading('4.6 SQL Explorer', level=2)
para(
    'The SQL Explorer tab allows users to execute arbitrary SELECT/INSERT/UPDATE queries against '
    'the database. Results are displayed in a dynamically generated table. Destructive operations '
    '(DROP, TRUNCATE, DELETE, ALTER, CREATE) are blocked at the API level. This tool is intended '
    'for testing and demonstration purposes.'
)

img_placeholder('Figure 4d: Frontend — SQL Explorer with Query Results')

# ══════════════════════════════════════════════════════════════
# 5. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════
heading('5. Technology Stack Summary', level=1)
add_table(
    ['Layer', 'Technology', 'Version'],
    [
        ['Database', 'PostgreSQL', '17'],
        ['Backend Framework', 'Django', '6.0.3'],
        ['Database Driver', 'psycopg2', '2.9.12'],
        ['Frontend', 'HTML5 / CSS3 / JavaScript (Vanilla)', '—'],
        ['Fonts', 'Plus Jakarta Sans, Space Grotesk, JetBrains Mono', '—'],
        ['Platform', 'Windows 11', '—'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. CONCLUSION
# ══════════════════════════════════════════════════════════════
heading('6. Conclusion', level=1)
para(
    'The Yogurt Processing System demonstrates a complete implementation of a database-driven '
    'web application with all business logic enforced at the PostgreSQL level. The three-tier '
    'architecture separates concerns cleanly:'
)
add_bullet('Database: All constraints, triggers, and referential integrity rules are defined and enforced in PostgreSQL.')
add_bullet('Backend: Django serves as a lightweight REST API layer executing raw SQL queries with no ORM abstraction.')
add_bullet('Frontend: A vanilla JavaScript SPA provides a clean, responsive user interface with real-time feedback.')

para(
    'All 13 API endpoints are fully functional and tested. Four security tests confirm that '
    'database-level triggers and constraints correctly block prohibited operations. The audit '
    'system captures all data changes across 8 tables with full JSONB snapshots.'
)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— End of Section —')
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)

# ── Save ──
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    'System_Implementation_Report.docx'
)
doc.save(output_path)
print(f'Report generated: {output_path}')
